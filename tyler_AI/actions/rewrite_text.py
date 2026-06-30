from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import logging
import re

logger = logging.getLogger(__name__)


class WordCreator:
    """Создатель форматированных Word документов из структурированного текста"""

    def run(self, vba_text: str, chat_id: str) -> dict:
        """
        Создаёт форматированный Word документ из структурированного текста
        
        Args:
            vba_text: Текст с маркерами форматирования в комментариях
            
        Returns:
            dict: Результат операции
        """
        try:
            # Получаем путь к папке Documents пользователя
            documents_path = os.path.expanduser("~/Documents")

            # Создаем подпапку для наших документов
            output_dir = os.path.join(documents_path, "generated_docs")
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
                logger.info(f"Создана папка: {output_dir}")

            # Генерируем имя файла
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}.docx"
            filepath = os.path.join(output_dir, filename)

            # Создаем документ
            doc = Document()

            # Устанавливаем стили по умолчанию
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

            # Обрабатываем строки
            lines = vba_text.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i].rstrip()

                # Пропускаем пустые строки с комментариями
                if not line or line == "'":
                    i += 1
                    continue

                # Убираем символ комментария
                if line.startswith("'"):
                    content = line[1:].lstrip()
                else:
                    content = line

                # Определяем тип форматирования по маркерам
                if "════════════════════════════════════════════════════════════════" in line:
                    # Заголовок h1
                    if i + 1 < len(lines):
                        title_line = lines[i + 1]
                        if title_line.startswith(
                                "'") and not "════════════════════════════════════════════════════════════════" in title_line:
                            title = title_line[1:].strip()
                            # Добавляем заголовок первого уровня
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run(title)
                            run.font.size = Pt(18)
                            run.font.bold = True
                            run.font.name = 'Arial'
                            p.space_before = Pt(12)
                            p.space_after = Pt(6)
                            i += 2  # Пропускаем обработанные строки
                            continue

                elif "────────────────────────────────────────────────────────────────" in line:
                    # Заголовок h2
                    if i + 1 < len(lines):
                        title_line = lines[i + 1]
                        if title_line.startswith(
                                "'") and not "────────────────────────────────────────────────────────────────" in title_line:
                            title = title_line[1:].strip()
                            # Добавляем заголовок второго уровня
                            p = doc.add_paragraph()
                            run = p.add_run(title)
                            run.font.size = Pt(14)
                            run.font.bold = True
                            run.font.name = 'Arial'
                            p.space_before = Pt(12)
                            p.space_after = Pt(6)
                            i += 2  # Пропускаем обработанные строки
                            continue

                elif content.startswith('•'):
                    # Заголовок h3
                    p = doc.add_paragraph()
                    run = p.add_run(content[1:].strip())
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.name = 'Arial'
                    p.space_before = Pt(8)
                    p.space_after = Pt(4)

                elif content.startswith('◦'):
                    # Заголовок h4
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    run = p.add_run(content[1:].strip())
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.italic = True
                    run.font.name = 'Arial'

                elif content.startswith('-'):
                    # Элемент списка
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(-0.25)
                    run = p.add_run('• ' + content[1:].strip())
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'

                elif content.strip() and not content.startswith(('Источник:', 'Дата:')):
                    # Обычный абзац
                    # Проверяем, не является ли это частью длинного текста
                    if content and len(content) > 0:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(8)

                        # Разбиваем длинные строки на предложения для лучшего форматирования
                        sentences = re.split(r'(?<=[.!?])\s+', content)
                        for sentence in sentences:
                            if sentence:
                                run = p.add_run(sentence + ' ')
                                run.font.size = Pt(11)
                                run.font.name = 'Arial'

                i += 1

            # Сохраняем документ
            doc.save(filepath)

            # Проверяем результат
            if os.path.exists(filepath) and os.path.getsize(filepath) > 0:
                logger.info(f"Документ сохранен: {filepath}")
                return {
                    "success": True,
                    "filename": filename,
                    "path": os.path.abspath(filepath)
                }
            else:
                return {
                    "success": False,
                    "error": "Файл не был создан или пустой"
                }

        except Exception as e:
            logger.error(f"Ошибка создания документа: {e}")
            return {
                "success": False,
                "error": str(e)
            }
