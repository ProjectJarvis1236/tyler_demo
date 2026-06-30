from .Web_parse import PageParser
import logging
import httpx
from docx import Document # Импортируем python-docx для создания .docx
from typing import Dict, Any, List

import configs

logger = logging.getLogger(__name__)

# Константы для API
API_URL = "https://apifreellm.com/api/v1/chat"
API_KEY = configs.APIFREELLM_KEY


class WordCoordinator:
    """Координатор создания Word документов из веб-страниц по списку ссылок"""

    def __init__(self):
        self.page_parser = PageParser()

    async def run(self, params: dict, chat_id: str) -> Dict[str, Any]:        
        """
        Основной метод для создания Word документа из списка ссылок.

        Args:
            params (dict): {
                "urls": ["https://...", "https://..."],  // Обязательный параметр
                "task_description": "Перепиши текст, структурируй, разбей на абзацы...", // Обязательный параметр
                "filename": "output.docx" // Опциональный параметр
            }

        Returns:
            Dict с результатом операции
        """

        print(f"[DEBUG WordCoordinator.run] Вызван с params: {params}")
        urls = params.get("urls")
        task_description = params.get("task_description")
        requested_filename = params.get("filename")

        if not urls or not isinstance(urls, list) or len(urls) == 0:
            return {
                "success": False,
                "reply": "Ошибка: В параметрах 'params' не найден или пустой список 'urls'."
            }

        if not task_description:
            return {
                "success": False,
                "reply": "Ошибка: В параметрах 'params' не найдено описание задачи 'task_description'."
            }

        # Проверка API ключа
        if not API_KEY:
             logger.error("API ключ для apifreellm не найден в configs.APIFREELLM_KEY")
             return {
                 "success": False,
                 "reply": "Ошибка: Не задан API ключ для LLM."
             }

        logger.info(f"Получен запрос на создание документа из {len(urls)} ссылок.")


        all_contents = []
        successful_parses = 0
        errors = []

        for url in urls:
            logger.info(f"Парсим: {url}")
            parse_result = self.page_parser.parse_single_url(url)
            if parse_result["success"]:
                all_contents.append({
                    "url": parse_result["url"],
                    "content": parse_result["content"]
                })
                successful_parses += 1
            else:
                errors.append(parse_result["error"])
                logger.warning(f"Ошибка парсинга {url}: {parse_result['error']}")

        if not all_contents:
            error_msg = f"Не удалось спарсить ни одну из {len(urls)} ссылок. Ошибки: {'; '.join(errors[:3])}..." if errors else "Неизвестная ошибка при парсинге."
            return {
                "success": False,
                "reply": error_msg
            }

        logger.info(f"Успешно спарсено {successful_parses} из {len(urls)} ссылок.")


        combined_content_parts = []
        for item in all_contents:
            combined_content_parts.append(f"--- ИСТОЧНИК: {item['url']} ---\n{item['content']}\n\n")
        combined_content = "".join(combined_content_parts)

  
        logger.info("Отправляем содержимое в LLM для обработки...")
        llm_result = await self._call_llm_for_rewrite(combined_content, task_description)

        if not llm_result or not llm_result.get("processed_text"):
             error_detail = llm_result.get('error', 'LLM не вернул текст') if llm_result else 'Неизвестная ошибка LLM'
             return {
                "success": False,
                "reply": f"Ошибка: LLM не вернул обработанный текст. Причина: {error_detail}"
            }

        processed_text = llm_result["processed_text"]
        # Используем запрошенное имя файла или генерируемое LLM
        final_filename = requested_filename or llm_result.get("generated_filename") or "generated_document.docx"


        logger.info(f"Создаём Word документ: {final_filename}")
        creation_result = self._create_word_document(processed_text, all_contents, final_filename)

        if not creation_result["success"]:
            return {
                "success": False,
                "reply": f"Ошибка при создании документа: {creation_result.get('error', 'Неизвестная ошибка')}"
            }


        return {
            "success": True,
            "reply": f"Документ '{final_filename}' успешно создан.",
            "result": {
                "filename": final_filename,
                "filepath": creation_result["filepath"],
                "source_count": len(all_contents),
                "processed_text_preview": processed_text[:200] + "..." # Предварительный просмотр
            }
        }

    async def _call_llm_for_rewrite(self, content: str, task_description: str) -> Dict[str, Any]:
        """
        Отправляет объединённый контент и описание задачи в NVIDIA LLM.
        Ожидает, что LLM вернёт обработанный текст и, возможно, название файла.
        """

        NVIDIA_API_KEY = configs.OPENROUTER_KEY
        NVIDIA_URL = "https://openrouter.ai/api/v1/chat/completions"
        if not NVIDIA_API_KEY:
            logger.error("NVIDIA API ключ не найден в configs.NVIDIA_API_KEY")
            print(f"[DEBUG LLM RESULT] Ошибка: NVIDIA API ключ не задан.")
            return {"success": False, "error": "NVIDIA API ключ не задан."}

        system_prompt = f"""
        Ты — опытный редактор и исследователь.
        Ниже представлен текст, объединённый из нескольких источников (разделённых метками "--- ИСТОЧНИК: ... ---").
        Твоя задача — внимательно прочитать его, выполнить следующие действия и вернуть ТОЛЬКО ИТОГОВЫЙ ТЕКСТ:

        ЗАДАНИЕ: {task_description}

        ТРЕБОВАНИЯ К ВЫХОДУ:
        1. Верни ТОЛЬКО обработанный текст. Не добавляй комментарии, пояснения, заголовки вроде "Результат:" или "Сгенерированный текст:".
        2. Если задание подразумевает генерацию названия документа, добавь его в начале текста в виде заголовка первого уровня (например, "# Название документа"), а затем основной текст.
        3. Не используй разметку Markdown (__, *, **, ``` и т.д.), если только она не входит в задание. Используй простой текст с абзацами (переводы строк).
        """

        user_prompt = f""" ИСХОДНЫЙ ТЕКСТ: {content} """

        payload = {
            "model": "nvidia/nemotron-3-super-120b-a12b:free", # Указываем модель
            "messages": [
                {
                    "role": "system",
                    "content": system_prompt.strip() # Убираем лишние пробелы/переводы строк
                },
                {
                    "role": "user",
                    "content": user_prompt.strip() # Убираем лишние пробелы/переводы строк
                }
            ],
            "temperature": 0.7, # Можно настроить
            "max_tokens": 4000, # Увеличьте при необходимости
            "stream": False # Пока без потоковой передачи
        }

        headers = {
            "Authorization": f"Bearer {NVIDIA_API_KEY}",
            "Content-Type": "application/json"
        }

        logger.debug(f"Отправка запроса в NVIDIA LLM, длина system + user промпта: {len(system_prompt) + len(user_prompt)}")
        print(f"[DEBUG LLM CALL] Отправляю POST запрос на {NVIDIA_URL}")
        print(f"[DEBUG LLM CALL] Payload: {payload}") # Выводит весь payload, может быть длинным

        try:
            async with httpx.AsyncClient(timeout=45) as client:
                r = await client.post(NVIDIA_URL, headers=headers, json=payload)

            logger.debug(f"Ответ NVIDIA LLM: статус {r.status_code}")
            print(f"[DEBUG LLM RESULT]   Статус: {r.status_code}")
            print(f"[DEBUG LLM RESULT]   Тело ответа (первые 500 символов): {r.text[:500]}...")

            if r.status_code == 200:
                data = r.json()
                # Проверяем структуру ответа NVIDIA
                choices = data.get("choices", [])
                if choices and len(choices) > 0:
                    message_content = choices[0].get("message", {}).get("content", "").strip()

                    if message_content:
                        # Попробуем извлечь название файла из заголовка, если он есть в ответе LLM
                        generated_filename = None
                        lines = message_content.split('\n')
                        if lines and lines[0].strip().startswith('# '):
                             generated_filename = lines[0].strip()[2:] + ".docx" # Убираем '# ' и добавляем .docx
                             message_content = '\n'.join(lines[1:]).strip() # Обновляем текст, убирая заголовок

                        print(f"[DEBUG LLM RESULT] Успешно получен ответ от NVIDIA LLM.")
                        return {
                            "success": True,
                            "processed_text": message_content,
                            "generated_filename": generated_filename
                        }
                    else:
                        logger.error("NVIDIA LLM вернул пустой ответ в choices[0].message.content.")
                        print(f"[DEBUG LLM RESULT] NVIDIA LLM вернул пустой ответ в choices[0].message.content.")
                        return {"success": False, "error": "NVIDIA LLM вернул пустой ответ."}

                else:
                    logger.error(f"NVIDIA LLM вернул 200 OK, но структура 'choices' некорректна или пуста: {data}")
                    print(f"[DEBUG LLM RESULT] NVIDIA LLM 200 OK, но 'choices' некорректны: {data}")
                    return {"success": False, "error": "NVIDIA LLM вернул некорректный ответ."}

            elif r.status_code == 429:
                 logger.warning("NVIDIA LLM API вернул 429 (Too Many Requests).")
                 print(f"[DEBUG LLM RESULT] NVIDIA LLM вернул 429 (Too Many Requests).")
                 return {"success": False, "error": "Превышено количество запросов к NVIDIA LLM. Попробуйте позже."}

            else:
                logger.error(f"NVIDIA LLM API ошибка {r.status_code}: {r.text[:200]}")
                print(f"[DEBUG LLM RESULT] NVIDIA LLM вернул ошибку {r.status_code}: {r.text[:200]}...")
                return {"success": False, "error": f"API ошибка {r.status_code}: {r.text[:100]}"}

        except httpx.ReadTimeout:
            logger.error("NVIDIA LLM API запрос превысил время ожидания.")
            print(f"[DEBUG LLM RESULT] Ошибка: Таймаут запроса к NVIDIA LLM ({45} сек).")
            return {"success": False, "error": "Таймаут запроса к NVIDIA LLM."}

        except Exception as e:
            logger.error(f"Ошибка при вызове NVIDIA LLM: {e}")
            print(f"[DEBUG LLM RESULT] Исключение при вызове NVIDIA LLM: {e}")
            return {"success": False, "error": str(e)}


    def _create_word_document(self, processed_text: str, sources: List[Dict[str, str]], filename: str) -> Dict[str, Any]:
        """
        Создаёт Word документ (.docx) с текстом и списком источников.
        """
        try:
            doc = Document()

            # Добавляем основной текст
            # Разбиваем на абзацы по переводам строк
            paragraphs = processed_text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip(): # Не добавляем пустые абзацы
                    doc.add_paragraph(para_text.strip())

            #doc.add_page_break() # Или просто абзац с заголовком
            doc.add_heading('Список источников', level=1)

            for index, source_item in enumerate(sources, start=1): # enumerate даёт (0, item), (1, item)... start=1 делает (1, item), (2, item)...
                doc.add_paragraph(f"{index}. {source_item['url']}", style='Normal') # Добавляем параграф с нумерованным текстом

            output_folder_path = configs.USER_FOLDER

            import os
            os.makedirs(output_folder_path, exist_ok=True)

            filepath = os.path.join(output_folder_path, filename)

            doc.save(filepath)

            logger.info(f"Word документ сохранён: {filepath}")
            return {
                "success": True,
                "filepath": filepath
            }

        except Exception as e:
            logger.error(f"Ошибка при создании Word документа: {e}")
            return {
                "success": False,
                "error": str(e)
            }
