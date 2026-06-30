import json
from uuid import uuid4
from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
import httpx

import configs


class DocxParser:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self.document_model = {"elements": []}
        self.render_map = {}

    @staticmethod
    def _iter_blocks(parent):
        if hasattr(parent, "element"):
            parent_elm = parent.element.body
        else:
            parent_elm = parent._tc

        for child in parent_elm.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, parent)
            elif child.tag.endswith("}tbl"):
                yield Table(child, parent)

    def _parse_run(self, run, paragraph_id, run_index):
        run_id = self._id("run")
        self.render_map[run_id] = {
            "type": "run",
            "paragraph_id": paragraph_id,
            "run_index": run_index
        }

        return {
            "id": run_id,
            "text": run.text,
            "style": {
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": str(run.font.size) if run.font.size else None,
            }
        }

    def _parse_paragraph(self, paragraph):
        paragraph_id = self._id("paragraph")
        data = {
            "id": paragraph_id,
            "type": "paragraph",
            "text": paragraph.text,
            "runs": []
        }

        self.render_map[paragraph_id] = {
            "type": "paragraph",
            "object": paragraph
        }

        for i, run in enumerate(paragraph.runs):
            data["runs"].append(self._parse_run(run, paragraph_id, i))
        return data

    def _parse_cell(self, cell: _Cell):
        cell_id = self._id("cell")
        data = {
            "id": cell_id,
            "type": "table_cell",
            "text": cell.text,
            "children": []
        }
        self.render_map[cell_id] = {
            "type": "cell",
            "object": cell
        }
        for block in self._iter_blocks(cell):
            if isinstance(block, Paragraph):
                data["children"].append(self._parse_paragraph(block))
            elif isinstance(block, Table):
                data["children"].append(self._parse_table(block))

        return data

    def _parse_table(self, table):
        table_id = self._id("table")
        data = {
            "id": table_id,
            "type": "table",
            "rows": []
        }
        self.render_map[table_id] = {
            "type": "table",
            "object": table
        }
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(self._parse_cell(cell))

            data["rows"].append(row_data)
        return data

    def parse(self):
        self.document_model = {"elements": []}

        for block in self._iter_blocks(self.doc):
            if isinstance(block, Paragraph):
                self.document_model["elements"].append(self._parse_paragraph(block))
            elif isinstance(block, Table):
                self.document_model["elements"].append(self._parse_table(block))

        return self.document_model

    def apply_operations(self, operations: list):
        for op in operations:
            op_type = op.get("type")
            target = op.get("target")
            content = op.get("content", "")

            if target not in self.render_map: continue

            target_data = self.render_map[target]

            if op_type == "replace_text":
                if target_data["type"] == "paragraph":
                    paragraph = target_data["object"]
                    if paragraph.runs:
                        paragraph.runs[0].text = content

                        for run in paragraph.runs[1:]:
                            run.text = ""

                    else:
                        paragraph.add_run(content)

            elif target_data["type"] == "cell":
                cell = target_data["object"]
                if cell.paragraphs:
                    p = cell.paragraphs[0]
                    if p.runs:
                        p.runs[0].text = content

                        for run in p.runs[1:]:
                            run.text = ""

                    else:
                        p.add_run(content)

    def save(self, output_path: str):
        self.doc.save(output_path)

    @staticmethod
    def _id(prefix: str):
        return f"{prefix}_{uuid4().hex[:8]}"

    def get_compressed_model(self):
        def compress(element):
            result = {
                "id": element["id"],
                "type": element["type"]
            }

            if "text" in element:
                result["text"] = element["text"]

            if element["type"] == "table":
                result["rows"] = [
                    [
                        {
                            "id": cell["id"],
                            "text": cell["text"]
                        }
                        for cell in row
                    ]
                    for row in element["rows"]
                ]

            return result

        return {
            "elements": [
                compress(el) for el in self.document_model["elements"]
            ]
        }


API_KEYS = {
    "nvidia/nemotron-3-super-120b-a12b:free": configs.OPENROUTER_KEY,
    "apifreellm": configs.APIFREELLM_KEY
}


class DocumentWorker:
    def __init__(self):
        pass

    async def call_llm(self, payload: dict):
        model = "nvidia/nemotron-3-super-120b-a12b:free"
        API_KEY = API_KEYS[model]
        url = configs.URLS[model]

        headers = {
            "Authorization": f"Bearer {API_KEY}",
            "Content-Type": "application/json"
        }

        system_prompt = SYSTEM_PROMPT

        data = {
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": json.dumps(payload, ensure_ascii=False)
                }
            ]
        }

        async with httpx.AsyncClient(timeout=configs.TIMEOUT) as client:
            response = await client.post(url, headers=headers, json=data)
            result = response.json()

        content = result["choices"][0]["message"]["content"]
        return json.loads(content)

    async def run(self, params: dict, chat_id: str):
        """
        params:
        {
            "file_path": "input.docx",
            "output_path": "output.docx",
            "prompt": "Заполни документ",
            "llm_state": {}
        }
        """

        file_path = params["file_path"]
        output_path = params.get("output_path", "output.docx")
        prompt = params["prompt"]
        llm_state = params.get("llm_state", {})

        parser = DocxParser(file_path)
        parser.parse()
        compressed_document = parser.get_compressed_model()

        llm_payload = {
            "prompt": prompt,
            "document": compressed_document,
            "state": llm_state
        }

        llm_response = await self.call_llm(llm_payload)
        operations = llm_response.get("operations", [])

        parser.apply_operations(operations)
        parser.save(output_path)

        return {
            "status": "success",
            "output_path": output_path,
            "llm_response": llm_response
        }


SYSTEM_PROMPT = """
Ты — AI для редактирования документов Word.

Тебе передаются:
1. prompt — запрос пользователя.
2. document — структура документа в JSON.
3. state — промежуточное состояние прошлых вызовов.

Твоя задача:
- проанализировать документ целиком;
- понять, какие элементы нужно изменить;
- вернуть только JSON с инструкциями для редактирования.

ВАЖНЫЕ ПРАВИЛА:
1. Никогда не переписывай документ целиком.
2. Изменяй только необходимые элементы.
3. Используй только существующие id элементов из document.
4. Не придумывай новые id.
5. Не добавляй пояснения, комментарии, markdown, текст до или после JSON.
6. Ответ должен быть только валидным JSON.
7. Не используй ```json и другие markdown блоки.
8. Ты должен **всегда** возвращать ответ со статусом "success" и списком операций.
   Даже если данных мало, додумай реалистичные значения, опираясь на контекст документа.
9. Ты переписываешь элемень целиком, внося необходимые изменения. Например, если будет написано "от (должность)" ты должен вернуть "от отдела продаж", а не просто "отдела продаж"

ДОСТУПНЫЕ ОПЕРАЦИИ:
- replace_text

Формат ответа (единственно допустимый):

{
  "status": "success",
  "operations": [
    {
      "type": "replace_text",
      "target": "id_элемента",
      "content": "новый текст"
    }
  ]
}

ПРАВИЛА ЗАПОЛНЕНИЯ:
- учитывай смысл всего документа, а не только отдельного блока;
- если в таблице одна ячейка содержит тему, а соседняя пустая, можно заполнить пустую ячейку содержимым по теме;
- не изменяй заполненные данные без необходимости;
- не изменяй заголовки, подписи, служебный текст, если пользователь явно не просил.
"""
